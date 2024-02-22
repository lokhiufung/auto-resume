from docx import Document
from docx.shared import Inches
import docx.shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_no_border(cell):
    """Set no borders for a cell."""
    tc_pr = cell._element.get_or_add_tcPr()  # Get the <w:tcPr> element, adding if not present
    # Remove any existing borders
    tcBorders = tc_pr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tc_pr.remove(tcBorders)
    
    # Create a new <w:tcBorders> element
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border}')
        border_el.set(qn('w:val'), 'nil')  # Set border value to 'nil' to make it invisible
        tcBorders.append(border_el)
    
    tc_pr.append(tcBorders)  # Add the new <w:tcBorders> to the cell properties


# Define a function to create the intro section
def add_intro(doc, resume):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Bold your title
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(resume['title'])
    run.bold = True  # Making the job title bold
    run.font.size = docx.shared.Pt(16)
    for info in [resume['location'], resume['phone'], resume['email'],
                 resume['linkedin'], resume.get('github', ''), resume.get('website', '')]:
        paragraph = cell.add_paragraph(info)
        paragraph.paragraph_format.space_after = docx.shared.Pt(0)
    set_no_border(cell)


def add_summary(doc, resume):
    # Create a single-cell table for the summary section
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add the 'Summary' heading to the cell
    cell.text = 'Summary'
    summary_heading = cell.paragraphs[0]
    summary_heading.style = doc.styles['Heading 1']

    # Add the actual summary text to the cell
    summary_paragraph = cell.add_paragraph(resume['summary'])
    summary_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
    
    # Remove the border from the cell
    set_no_border(cell)

# Define a function to create the experience section
def add_experience(doc, resume):
    doc.add_heading('Experience', level=1)
    for job in resume['jobs']:
        # Create a table for each job entry
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Inches(4)
        table.columns[1].width = Inches(2)
         # Set the table to have no borders
        for row in table.rows:
            for cell in row.cells:
                set_no_border(cell)

        # Job title
        cell = table.cell(0, 0)
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(resume['title'])
        run.bold = True  # Making the job title bold
        cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
        
        # Company name and location in the left cell
        cell = table.cell(1, 0)
        cell.text = f"{job['company']}, {job['location']}"
        # cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)

        # Duration in the right cell, aligned right
        cell = table.cell(1, 1)
        cell.text = job['duration']
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # cell.paragraphs[0].runs[0].bold = True

        # Job Descriptions in the subsequent rows of the table
        # Add a new row for bullet points
        row = table.add_row()
        desc_cell = row.cells[0]
        desc_cell.merge(row.cells[1])
        for item in job['desc']:
            # Merge cells for the description to span both columns
            desc_paragraph = desc_cell.add_paragraph(item, style='ListBullet')
            # Ensure no space after the paragraph for alignment
            desc_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
        set_no_border(desc_cell)

        # set_no_border(table.rows[0].cells[0])
        # set_no_border(table.rows[0].cells[1])


# Define a function to create the education section
def add_education(doc, resume):
    # ... (code to create the education section)
    ...

# Define a function to create the skills section
def add_skills(doc, resume):
    # ... (code to create the skills section)
    ...



def create_resume(resume):
    doc = Document()
    # Set the margins for the document
    section = doc.sections[0]
    section.top_margin = Inches(0.5)   # top margin
    section.bottom_margin = Inches(0.5)  # bottom margin
    section.left_margin = Inches(0.5)   # left margin
    section.right_margin = Inches(0.5)  # right margin

    # Title (Name)
    doc.add_heading(resume['name'], level=0)

    # Add Intro
    add_intro(doc, resume)
    # Add Summary
    add_summary(doc, resume)
    # Add Experience
    add_experience(doc, resume)
    # Add Education
    add_education(doc, resume)
    # Add Skills
    add_skills(doc, resume)
    
    return doc