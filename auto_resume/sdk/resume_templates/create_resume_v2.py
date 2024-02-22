from docx import Document
from docx.shared import Inches
import docx.shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_no_border(cell):
    """Set no borders for a cell."""
    tc_pr = cell._element.get_or_add_tcPr()  # Get the <w:tcPr> element, adding if not present
    # Remove any existing borders
    tcBorders = tc_pr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tc_pr.remove(tcBorders)
    
    # Create a new <w:tcBorders> element
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border}')
        border_el.set(qn('w:val'), 'nil')  # Set border value to 'nil' to make it invisible
        tcBorders.append(border_el)
    
    tc_pr.append(tcBorders)  # Add the new <w:tcBorders> to the cell properties


# Define a function to create the intro section
def add_intro(doc, resume):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    for info in [resume['title'], resume['location'], resume['phone'], resume['email'],
                 resume['linkedin'], resume.get('github', ''), resume.get('website', '')]:
        paragraph = cell.add_paragraph(info)
        paragraph.paragraph_format.space_after = docx.shared.Pt(0)
    set_no_border(cell)


# Define a function to create the summary section
def add_summary(doc, resume):
    doc.add_heading('Summary', level=1)
    summary_paragraph = doc.add_paragraph(resume['summary'])
    summary_paragraph.paragraph_format.space_after = docx.shared.Pt(0)


# Define a function to create the experience section
def add_experience(doc, resume):
    doc.add_heading('Experience', level=1)
    for job in resume['jobs']:
        # Create a table for each job entry
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Inches(4)
        table.columns[1].width = Inches(2)
         # Set the table to have no borders
        for row in table.rows:
            for cell in row.cells:
                set_no_border(cell)

        # Company name and location in the left cell
        cell = table.cell(0, 0)
        cell.text = f"{job['company']}, {job['location']}"
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)

        # Duration in the right cell, aligned right
        cell = table.cell(0, 1)
        cell.text = job['duration']
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].runs[0].bold = True

        # Job Descriptions in the subsequent rows of the table
        for item in job['desc']:
            # Add a new row for each bullet point
            row = table.add_row()
            # Merge cells for the description to span both columns
            desc_cell = row.cells[0]
            desc_cell.merge(row.cells[1])
            desc_paragraph = desc_cell.add_paragraph(item, style='ListBullet')
            # Ensure no space after the paragraph for alignment
            desc_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
            set_no_border(desc_cell)

        # set_no_border(table.rows[0].cells[0])
        # set_no_border(table.rows[0].cells[1])


# Define a function to create the education section
def add_education(doc, resume):
    # ... (code to create the education section)
    ...

# Define a function to create the skills section
def add_skills(doc, resume):
    # ... (code to create the skills section)
    ...


def create_resume(resume):
    doc = Document()
    # Set the margins for the document
    section = doc.sections[0]
    section.top_margin = Inches(0.5)   # top margin
    section.bottom_margin = Inches(0.5)  # bottom margin
    section.left_margin = Inches(0.5)   # left margin
    section.right_margin = Inches(0.5)  # right margin

    # Title (Name)
    doc.add_heading(resume['name'], level=0)

    # Contact details
    p = doc.add_paragraph()
    run = p.add_run(resume['title'])
    run.bold = True  # Making the job title bold
    run.font.size = docx.shared.Pt(16)

    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['location'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['phone'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['email'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['linkedin'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    if 'github' in resume:
        p = doc.add_paragraph(resume['github'])
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    if 'website' in resume:
        p = doc.add_paragraph(resume['website'])
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    # Summary
    doc.add_heading('Summary', level=1)
    summary_text = resume['summary']
    doc.add_paragraph(summary_text)

    
    # Experience
    doc.add_heading('Experience', level=1)

    for job in resume['jobs']:
        # Job title
        p = doc.add_paragraph()
        run = p.add_run(job['title'])
        run.bold = True
        p.paragraph_format.line_spacing = docx.shared.Pt(12)  # Adjust line spacing to 12 points
        p.paragraph_format.space_after = docx.shared.Pt(0)

        # Table for Company, Location, and Duration on the same line
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'  # Apply the 'Table Grid' style

        # Set widths for consistency
        table.columns[0].width = Inches(4)
        table.columns[1].width = Inches(2)

        # Set the table to have no borders
        for row in table.rows:
            for cell in row.cells:
                set_no_border(cell)

        # Company name and location in the left cell
        cell = table.cell(0, 0)
        cell.text = f"{job['company']}, {job['location']}"
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)

        # Duration in the right cell, aligned right
        cell = table.cell(0, 1)
        cell.text = job['duration']
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].runs[0].bold = True

        # Job Descriptions in the subsequent rows of the table
        for item in job['desc']:
            # Add a new row for each bullet point
            row = table.add_row()
            # Merge cells for the description to span both columns
            desc_cell = row.cells[0]
            desc_cell.merge(row.cells[1])
            desc_paragraph = desc_cell.add_paragraph(item, style='ListBullet')
            # Ensure no space after the paragraph for alignment
            desc_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
            set_no_border(desc_cell)

    # Education
    doc.add_heading('Education', level=1)
    for education in resume['educations']:
        p = doc.add_paragraph()
        run = p.add_run('{}, {}'.format(education['name'],education['location']))
        run.bold = True
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

        p = doc.add_paragraph(education['duration'])
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

        doc.add_paragraph(education['degree'])


    # Skills
    doc.add_heading('Skills', level=1)
    skills = resume['skills']
    skills = resume['skills']
    for skill in skills:
        p = doc.add_paragraph(skill, style='ListBullet')
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

    # Save the doc
    # doc.save('resume_template.docx')
    return doc

