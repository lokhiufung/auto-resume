from docx import Document
from docx.shared import Inches
import docx.shared

# use a table to contain all the work experiences

def create_resume(resume):
    doc = Document()
    # Set the margins for the document
    section = doc.sections[0]
    section.top_margin = Inches(0.5)   # top margin
    section.bottom_margin = Inches(0.5)  # bottom margin
    section.left_margin = Inches(0.5)   # left margin
    section.right_margin = Inches(0.5)  # right margin

    # Title (Name)
    doc.add_heading(resume['name'], level=0)

    # Contact details
    p = doc.add_paragraph()
    run = p.add_run(resume['title'])
    run.bold = True  # Making the job title bold
    run.font.size = docx.shared.Pt(16)

    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['location'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['phone'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['email'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    p = doc.add_paragraph(resume['linkedin'])
    p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    if 'github' in resume:
        p = doc.add_paragraph(resume['github'])
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    if 'website' in resume:
        p = doc.add_paragraph(resume['website'])
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
    # Summary
    doc.add_heading('Summary', level=1)
    summary_text = resume['summary']
    doc.add_paragraph(summary_text)

    
    # Experience
    doc.add_heading('Experience', level=1)

    for job in resume['jobs']:
        # Job title
        p = doc.add_paragraph()
        run = p.add_run(job['title'])
        run.bold = True  # Making the job title bold
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

        # Company and Location
        company_location = f"{job['company']}, {job['location']}"
        p = doc.add_paragraph(company_location)
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

        # Date/Duration
        p = doc.add_paragraph(job['duration'])

        # Job Description
        for item in job['desc']:
            p = doc.add_paragraph(item, style='ListBullet')
            p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points
        
        doc.add_paragraph()

    # Education
    doc.add_heading('Education', level=1)
    for education in resume['educations']:
        p = doc.add_paragraph()
        run = p.add_run('{}, {}'.format(education['name'],education['location']))
        run.bold = True
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

        p = doc.add_paragraph(education['duration'])
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

        doc.add_paragraph(education['degree'])


    # Skills
    doc.add_heading('Skills', level=1)
    skills = resume['skills']
    skills = resume['skills']
    for skill in skills:
        p = doc.add_paragraph(skill, style='ListBullet')
        p.paragraph_format.space_after = docx.shared.Pt(0)  # Set space after the paragraph to 0 points

    # Save the doc
    # doc.save('resume_template.docx')
    return doc

